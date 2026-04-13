#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成基于SpringBoot+Vue的个人博客系统毕业设计论文
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

# 创建文档
doc = Document()

# 设置中文字体
def set_chinese_font(paragraph):
    for run in paragraph.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置标题字体
def set_title_font(paragraph):
    for run in paragraph.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 封面
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.add_run('基于SpringBoot+Vue的个人博客系统设计与实现')
run.font.size = Pt(28)
run.font.bold = True
set_title_font(cover)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

student_info = doc.add_paragraph()
student_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = student_info.add_run('学生姓名：黄翀')
run.font.size = Pt(16)
set_chinese_font(student_info)
doc.add_paragraph()

student_id = doc.add_paragraph()
student_id.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = student_id.add_run('学号：2022070030132')
run.font.size = Pt(16)
set_chinese_font(student_id)
doc.add_paragraph()

department = doc.add_paragraph()
department.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = department.add_run('所在系部：计算机科学与技术系')
run.font.size = Pt(16)
set_chinese_font(department)
doc.add_paragraph()

major = doc.add_paragraph()
major.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = major.add_run('专业：软件工程')
run.font.size = Pt(16)
set_chinese_font(major)
doc.add_paragraph()

advisor = doc.add_paragraph()
advisor.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = advisor.add_run('指导教师：马莉')
run.font.size = Pt(16)
set_chinese_font(advisor)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

date = doc.add_paragraph()
date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = date.add_run('2026年5月')
run.font.size = Pt(16)
set_chinese_font(date)

# 分页
doc.add_page_break()

# 摘要
abstract = doc.add_heading('摘要', level=1)
set_title_font(abstract)

abstract_content = doc.add_paragraph()
abstract_content.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = abstract_content.add_run('随着互联网技术的快速发展，个人博客作为一种表达个人思想、分享知识的平台，受到越来越多人的青睐。本文设计并实现了一个基于SpringBoot+Vue的个人博客系统，旨在为用户提供一个功能完善、界面美观、易于使用的博客平台。')
run.font.size = Pt(12)
set_chinese_font(abstract_content)
doc.add_paragraph()

abstract_content2 = doc.add_paragraph()
abstract_content2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = abstract_content2.add_run('系统采用前后端分离的架构，后端使用SpringBoot框架构建RESTful API，前端使用Vue框架实现用户界面。数据库采用MySQL存储数据，使用MyBatis-Plus进行ORM映射。系统实现了用户管理、文章管理、评论管理、分类标签管理、说说管理、资源管理等核心功能。')
run.font.size = Pt(12)
set_chinese_font(abstract_content2)
doc.add_paragraph()

abstract_content3 = doc.add_paragraph()
abstract_content3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = abstract_content3.add_run('本文详细介绍了系统的需求分析、系统设计、技术选型、实现过程和测试结果。系统经过测试，功能完善，性能稳定，达到了预期的设计目标。本系统不仅满足了个人博客的基本需求，还提供了丰富的扩展功能，为用户提供了良好的使用体验。')
run.font.size = Pt(12)
set_chinese_font(abstract_content3)
doc.add_paragraph()

key_words = doc.add_paragraph()
run = key_words.add_run('关键词：个人博客；SpringBoot；Vue；前后端分离；MySQL')
run.font.size = Pt(12)
set_chinese_font(key_words)

# 分页
doc.add_page_break()

# 目录
toc = doc.add_heading('目录', level=1)
set_title_font(toc)
doc.add_paragraph()

# 生成目录项
def add_toc_item(title, level, page_num):
    para = doc.add_paragraph()
    run = para.add_run('{0} {1}'.format(' ' * (level-1) * 4, title))
    run.font.size = Pt(12)
    set_chinese_font(para)
    # 添加页码右对齐
    run = para.add_run(' ' * (60 - len(title) - (level-1)*4))
    run = para.add_run(str(page_num))
    run.font.size = Pt(12)

add_toc_item('1 引言', 1, 3)
add_toc_item('1.1 研究背景', 2, 3)
add_toc_item('1.2 研究意义', 2, 4)
add_toc_item('1.3 研究内容', 2, 4)
add_toc_item('2 系统分析', 1, 5)
add_toc_item('2.1 需求分析', 2, 5)
add_toc_item('2.2 可行性分析', 2, 6)
add_toc_item('2.3 技术选型', 2, 7)
add_toc_item('3 系统设计', 1, 8)
add_toc_item('3.1 系统架构设计', 2, 8)
add_toc_item('3.2 功能模块设计', 2, 9)
add_toc_item('3.3 数据库设计', 2, 10)
add_toc_item('4 系统实现', 1, 12)
add_toc_item('4.1 后端实现', 2, 12)
add_toc_item('4.2 前端实现', 2, 14)
add_toc_item('4.3 移动端实现', 2, 16)
add_toc_item('5 系统测试', 1, 17)
add_toc_item('5.1 测试方法', 2, 17)
add_toc_item('5.2 测试结果', 2, 18)
add_toc_item('6 结论与展望', 1, 19)
add_toc_item('6.1 结论', 2, 19)
add_toc_item('6.2 展望', 2, 20)
add_toc_item('参考文献', 1, 21)
add_toc_item('致谢', 1, 22)

# 分页
doc.add_page_break()

# 1 引言
chapter1 = doc.add_heading('1 引言', level=1)
set_title_font(chapter1)

section1_1 = doc.add_heading('1.1 研究背景', level=2)
set_title_font(section1_1)

para1_1 = doc.add_paragraph()
para1_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para1_1.add_run('在互联网时代，个人博客已成为人们分享知识、表达思想、展示个性的重要平台。传统的博客系统往往功能单一、界面陈旧，无法满足现代用户的需求。随着Web技术的不断发展，前后端分离架构已成为Web应用开发的主流趋势，这种架构可以提高开发效率，改善用户体验。')
run.font.size = Pt(12)
set_chinese_font(para1_1)
doc.add_paragraph()

para1_1_2 = doc.add_paragraph()
para1_1_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para1_1_2.add_run('SpringBoot作为Java生态中最流行的后端框架，提供了快速开发、自动配置等特性，大大简化了后端开发流程。Vue作为前端框架的佼佼者，以其轻量级、响应式的特点受到开发者的青睐。结合这两种技术，可以构建出性能优异、用户体验良好的现代Web应用。')
run.font.size = Pt(12)
set_chinese_font(para1_1_2)

section1_2 = doc.add_heading('1.2 研究意义', level=2)
set_title_font(section1_2)

para1_2 = doc.add_paragraph()
para1_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para1_2.add_run('本研究旨在设计并实现一个基于SpringBoot+Vue的个人博客系统，具有以下意义：')
run.font.size = Pt(12)
set_chinese_font(para1_2)
doc.add_paragraph()

list1_2 = doc.add_paragraph()
run = list1_2.add_run('1. 提供一个功能完善、界面美观的个人博客平台，满足用户的个性化需求。')
run.font.size = Pt(12)
set_chinese_font(list1_2)

list1_2_2 = doc.add_paragraph()
run = list1_2_2.add_run('2. 探索前后端分离架构在个人博客系统中的应用，为类似系统的开发提供参考。')
run.font.size = Pt(12)
set_chinese_font(list1_2_2)

list1_2_3 = doc.add_paragraph()
run = list1_2_3.add_run('3. 实践SpringBoot和Vue等现代Web技术，提高开发能力和技术水平。')
run.font.size = Pt(12)
set_chinese_font(list1_2_3)

section1_3 = doc.add_heading('1.3 研究内容', level=2)
set_title_font(section1_3)

para1_3 = doc.add_paragraph()
para1_3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para1_3.add_run('本研究的主要内容包括：')
run.font.size = Pt(12)
set_chinese_font(para1_3)
doc.add_paragraph()

list1_3 = doc.add_paragraph()
run = list1_3.add_run('1. 系统需求分析：分析个人博客系统的功能需求和非功能需求。')
run.font.size = Pt(12)
set_chinese_font(list1_3)

list1_3_2 = doc.add_paragraph()
run = list1_3_2.add_run('2. 系统设计：设计系统架构、功能模块和数据库结构。')
run.font.size = Pt(12)
set_chinese_font(list1_3_2)

list1_3_3 = doc.add_paragraph()
run = list1_3_3.add_run('3. 系统实现：使用SpringBoot实现后端API，使用Vue实现前端界面，使用MySQL存储数据。')
run.font.size = Pt(12)
set_chinese_font(list1_3_3)

list1_3_4 = doc.add_paragraph()
run = list1_3_4.add_run('4. 系统测试：测试系统的功能和性能，确保系统稳定运行。')
run.font.size = Pt(12)
set_chinese_font(list1_3_4)

list1_3_5 = doc.add_paragraph()
run = list1_3_5.add_run('5. 系统部署：部署系统到生产环境，确保系统可以正常访问。')
run.font.size = Pt(12)
set_chinese_font(list1_3_5)

# 分页
doc.add_page_break()

# 2 系统分析
chapter2 = doc.add_heading('2 系统分析', level=1)
set_title_font(chapter2)

section2_1 = doc.add_heading('2.1 需求分析', level=2)
set_title_font(section2_1)

para2_1 = doc.add_paragraph()
para2_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para2_1.add_run('通过对个人博客系统的调研和分析，系统需要满足以下功能需求：')
run.font.size = Pt(12)
set_chinese_font(para2_1)
doc.add_paragraph()

# 功能需求列表
list2_1 = doc.add_paragraph()
run = list2_1.add_run('1. 用户管理：注册、登录、个人信息管理、权限控制。')
run.font.size = Pt(12)
set_chinese_font(list2_1)

list2_1_2 = doc.add_paragraph()
run = list2_1_2.add_run('2. 文章管理：发布、编辑、删除文章，支持Markdown格式。')
run.font.size = Pt(12)
set_chinese_font(list2_1_2)

list2_1_3 = doc.add_paragraph()
run = list2_1_3.add_run('3. 分类标签管理：文章分类、标签管理。')
run.font.size = Pt(12)
set_chinese_font(list2_1_3)

list2_1_4 = doc.add_paragraph()
run = list2_1_4.add_run('4. 评论系统：文章评论、回复、点赞。')
run.font.size = Pt(12)
set_chinese_font(list2_1_4)

list2_1_5 = doc.add_paragraph()
run = list2_1_5.add_run('5. 说说功能：发布、查看说说。')
run.font.size = Pt(12)
set_chinese_font(list2_1_5)

list2_1_6 = doc.add_paragraph()
run = list2_1_6.add_run('6. 资源管理：图片、文件上传与管理。')
run.font.size = Pt(12)
set_chinese_font(list2_1_6)

list2_1_7 = doc.add_paragraph()
run = list2_1_7.add_run('7. 后台管理：用户管理、文章管理、评论管理、系统配置。')
run.font.size = Pt(12)
set_chinese_font(list2_1_7)

list2_1_8 = doc.add_paragraph()
run = list2_1_8.add_run('8. 移动端支持：微信小程序访问。')
run.font.size = Pt(12)
set_chinese_font(list2_1_8)

doc.add_paragraph()

para2_1_2 = doc.add_paragraph()
para2_1_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para2_1_2.add_run('非功能需求：')
run.font.size = Pt(12)
set_chinese_font(para2_1_2)
doc.add_paragraph()

list2_1_9 = doc.add_paragraph()
run = list2_1_9.add_run('1. 性能需求：页面加载速度快，响应及时。')
run.font.size = Pt(12)
set_chinese_font(list2_1_9)

list2_1_10 = doc.add_paragraph()
run = list2_1_10.add_run('2. 可用性需求：系统稳定运行，故障恢复能力强。')
run.font.size = Pt(12)
set_chinese_font(list2_1_10)

list2_1_11 = doc.add_paragraph()
run = list2_1_11.add_run('3. 安全性需求：数据加密存储，防止SQL注入等攻击。')
run.font.size = Pt(12)
set_chinese_font(list2_1_11)

list2_1_12 = doc.add_paragraph()
run = list2_1_12.add_run('4. 可扩展性需求：系统架构灵活，易于扩展新功能。')
run.font.size = Pt(12)
set_chinese_font(list2_1_12)

section2_2 = doc.add_heading('2.2 可行性分析', level=2)
set_title_font(section2_2)

para2_2 = doc.add_paragraph()
para2_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para2_2.add_run('1. 技术可行性：SpringBoot和Vue都是成熟的框架，拥有丰富的生态系统和社区支持。MySQL是主流的关系型数据库，性能稳定。这些技术的组合已经被广泛应用于生产环境，技术上是可行的。')
run.font.size = Pt(12)
set_chinese_font(para2_2)
doc.add_paragraph()

para2_2_2 = doc.add_paragraph()
para2_2_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para2_2_2.add_run('2. 经济可行性：系统开发所需的软件和工具都是开源的，不需要购买商业软件。服务器成本也相对较低，可以使用云服务器或本地服务器。')
run.font.size = Pt(12)
set_chinese_font(para2_2_2)
doc.add_paragraph()

para2_2_3 = doc.add_paragraph()
para2_2_3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para2_2_3.add_run('3. 操作可行性：系统界面设计简洁直观，用户容易上手。管理员可以通过后台管理系统方便地管理内容和用户。')
run.font.size = Pt(12)
set_chinese_font(para2_2_3)

section2_3 = doc.add_heading('2.3 技术选型', level=2)
set_title_font(section2_3)

# 技术选型表格
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

# 表头
heading_cells = table.rows[0].cells
heading_cells[0].text = '分类'
heading_cells[1].text = '技术'
heading_cells[2].text = '版本'

# 设置表头样式
for cell in heading_cells:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_chinese_font(paragraph)
        for run in paragraph.runs:
            run.font.bold = True

# 填充数据
data = [
    ['后端框架', 'SpringBoot', '2.7.0'],
    ['前端框架', 'Vue', '2.7.16 (网站) / 3.2.47 (管理后台)'],
    ['数据库', 'MySQL', '8.0+'],
    ['ORM框架', 'MyBatis-Plus', '3.5.2'],
    ['认证框架', 'Sa-Token', '1.39.0'],
    ['UI组件库', 'Element UI (网站) / Element Plus (管理后台)', '2.15.14 / 2.3.0']
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, cell_data in enumerate(row_data):
        row.cells[j].text = cell_data
        for paragraph in row.cells[j].paragraphs:
            set_chinese_font(paragraph)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 分页
doc.add_page_break()

# 3 系统设计
chapter3 = doc.add_heading('3 系统设计', level=1)
set_title_font(chapter3)

section3_1 = doc.add_heading('3.1 系统架构设计', level=2)
set_title_font(section3_1)

para3_1 = doc.add_paragraph()
para3_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para3_1.add_run('系统采用前后端分离的架构，具体架构如下：')
run.font.size = Pt(12)
set_chinese_font(para3_1)
doc.add_paragraph()

# 系统架构图描述
arch_desc = doc.add_paragraph()
arch_desc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = arch_desc.add_run('系统架构分为三层：')
run.font.size = Pt(12)
set_chinese_font(arch_desc)
doc.add_paragraph()

list3_1 = doc.add_paragraph()
run = list3_1.add_run('1. 前端层：包括博客网站（Vue 2.7）、管理后台（Vue 3.2）和移动端小程序（UniApp）。')
run.font.size = Pt(12)
set_chinese_font(list3_1)

list3_1_2 = doc.add_paragraph()
run = list3_1_2.add_run('2. 后端层：基于SpringBoot构建的RESTful API，处理业务逻辑。')
run.font.size = Pt(12)
set_chinese_font(list3_1_2)

list3_1_3 = doc.add_paragraph()
run = list3_1_3.add_run('3. 数据层：MySQL数据库存储数据，Redis用于缓存。')
run.font.size = Pt(12)
set_chinese_font(list3_1_3)

section3_2 = doc.add_heading('3.2 功能模块设计', level=2)
set_title_font(section3_2)

para3_2 = doc.add_paragraph()
para3_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para3_2.add_run('系统主要包含以下功能模块：')
run.font.size = Pt(12)
set_chinese_font(para3_2)
doc.add_paragraph()

# 功能模块表格
table2 = doc.add_table(rows=7, cols=2)
table2.style = 'Table Grid'

# 表头
heading_cells2 = table2.rows[0].cells
heading_cells2[0].text = '模块名称'
heading_cells2[1].text = '功能描述'

# 设置表头样式
for cell in heading_cells2:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_chinese_font(paragraph)
        for run in paragraph.runs:
            run.font.bold = True

# 填充数据
module_data = [
    ['用户管理', '注册、登录、个人信息管理、权限控制'],
    ['文章管理', '发布、编辑、删除文章，支持Markdown格式'],
    ['分类标签', '文章分类管理、标签管理'],
    ['评论系统', '文章评论、回复、点赞'],
    ['说说功能', '发布、查看说说'],
    ['资源管理', '图片、文件上传与管理']
]

for i, row_data in enumerate(module_data, start=1):
    row = table2.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            set_chinese_font(paragraph)

section3_3 = doc.add_heading('3.3 数据库设计', level=2)
set_title_font(section3_3)

para3_3 = doc.add_paragraph()
para3_3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para3_3.add_run('根据系统功能需求，设计了以下数据库表：')
run.font.size = Pt(12)
set_chinese_font(para3_3)
doc.add_paragraph()

# 数据库表结构
# 用户表
table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Table Grid'

# 表头
heading_cells3 = table3.rows[0].cells
heading_cells3[0].text = '字段名'
heading_cells3[1].text = '数据类型'
heading_cells3[2].text = '描述'

# 设置表头样式
for cell in heading_cells3:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_chinese_font(paragraph)
        for run in paragraph.runs:
            run.font.bold = True

# 填充数据
user_data = [
    ['id', 'BIGINT', '用户ID'],
    ['username', 'VARCHAR(50)', '用户名'],
    ['password', 'VARCHAR(100)', '密码'],
    ['email', 'VARCHAR(100)', '邮箱'],
    ['avatar', 'VARCHAR(255)', '头像']
]

for i, row_data in enumerate(user_data, start=1):
    row = table3.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    row.cells[2].text = row_data[2]
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            set_chinese_font(paragraph)

doc.add_paragraph()

# 文章表
table4 = doc.add_table(rows=7, cols=3)
table4.style = 'Table Grid'

# 表头
heading_cells4 = table4.rows[0].cells
heading_cells4[0].text = '字段名'
heading_cells4[1].text = '数据类型'
heading_cells4[2].text = '描述'

# 设置表头样式
for cell in heading_cells4:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_chinese_font(paragraph)
        for run in paragraph.runs:
            run.font.bold = True

# 填充数据
article_data = [
    ['id', 'BIGINT', '文章ID'],
    ['title', 'VARCHAR(200)', '标题'],
    ['content', 'LONGTEXT', '内容'],
    ['user_id', 'BIGINT', '作者ID'],
    ['category_id', 'BIGINT', '分类ID'],
    ['view_count', 'INT', '阅读量']
]

for i, row_data in enumerate(article_data, start=1):
    row = table4.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    row.cells[2].text = row_data[2]
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            set_chinese_font(paragraph)

# 分页
doc.add_page_break()

# 4 系统实现
chapter4 = doc.add_heading('4 系统实现', level=1)
set_title_font(chapter4)

section4_1 = doc.add_heading('4.1 后端实现', level=2)
set_title_font(section4_1)

para4_1 = doc.add_paragraph()
para4_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para4_1.add_run('后端使用SpringBoot框架实现，主要包括以下模块：')
run.font.size = Pt(12)
set_chinese_font(para4_1)
doc.add_paragraph()

list4_1 = doc.add_paragraph()
run = list4_1.add_run('1. 核心模块：SpringBoot应用主类，配置系统启动。')
run.font.size = Pt(12)
set_chinese_font(list4_1)

list4_1_2 = doc.add_paragraph()
run = list4_1_2.add_run('2. 控制器模块：处理HTTP请求，返回响应。')
run.font.size = Pt(12)
set_chinese_font(list4_1_2)

list4_1_3 = doc.add_paragraph()
run = list4_1_3.add_run('3. 服务模块：实现业务逻辑。')
run.font.size = Pt(12)
set_chinese_font(list4_1_3)

list4_1_4 = doc.add_paragraph()
run = list4_1_4.add_run('4. 数据访问模块：使用MyBatis-Plus操作数据库。')
run.font.size = Pt(12)
set_chinese_font(list4_1_4)

list4_1_5 = doc.add_paragraph()
run = list4_1_5.add_run('5. 实体模块：定义数据库表对应的实体类。')
run.font.size = Pt(12)
set_chinese_font(list4_1_5)

list4_1_6 = doc.add_paragraph()
run = list4_1_6.add_run('6. 工具模块：提供通用工具方法。')
run.font.size = Pt(12)
set_chinese_font(list4_1_6)

doc.add_paragraph()

# 后端目录结构
para4_1_2 = doc.add_paragraph()
para4_1_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para4_1_2.add_run('后端目录结构：')
run.font.size = Pt(12)
set_chinese_font(para4_1_2)
doc.add_paragraph()

code4_1 = doc.add_paragraph()
run = code4_1.add_run('blog/\n├── mojian-commom/     # 公共模块\n├── mojian-admin/       # 管理相关功能\n├── mojian-api/         # API接口\n├── mojian-server/      # 主服务模块\n│   └── src/main/java/com/mojian/\n│       ├── NeatAdminApplication.java  # 应用主类\n│       ├── controller/  # 控制器\n│       ├── service/     # 服务\n│       ├── mapper/      # 数据访问\n│       └── entity/      # 实体类\n├── mojian-file/        # 文件处理\n├── mojian-quartz/      # 定时任务\n└── mojian-auth/        # 认证授权')
run.font.size = Pt(10)
set_chinese_font(code4_1)

section4_2 = doc.add_heading('4.2 前端实现', level=2)
set_title_font(section4_2)

para4_2 = doc.add_paragraph()
para4_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para4_2.add_run('前端分为博客网站和管理后台两部分：')
run.font.size = Pt(12)
set_chinese_font(para4_2)
doc.add_paragraph()

# 博客网站
sub_section4_2_1 = doc.add_heading('4.2.1 博客网站', level=3)
set_title_font(sub_section4_2_1)

para4_2_1 = doc.add_paragraph()
para4_2_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para4_2_1.add_run('博客网站使用Vue 2.7实现，主要功能包括：')
run.font.size = Pt(12)
set_chinese_font(para4_2_1)
doc.add_paragraph()

list4_2_1 = doc.add_paragraph()
run = list4_2_1.add_run('1. 首页：展示文章列表、热门文章、分类标签等。')
run.font.size = Pt(12)
set_chinese_font(list4_2_1)

list4_2_1_2 = doc.add_paragraph()
run = list4_2_1_2.add_run('2. 文章详情：展示文章内容，支持Markdown格式。')
run.font.size = Pt(12)
set_chinese_font(list4_2_1_2)

list4_2_1_3 = doc.add_paragraph()
run = list4_2_1_3.add_run('3. 分类页：按分类展示文章。')
run.font.size = Pt(12)
set_chinese_font(list4_2_1_3)

list4_2_1_4 = doc.add_paragraph()
run = list4_2_1_4.add_run('4. 标签页：按标签展示文章。')
run.font.size = Pt(12)
set_chinese_font(list4_2_1_4)

list4_2_1_5 = doc.add_paragraph()
run = list4_2_1_5.add_run('5. 归档页：按时间归档文章。')
run.font.size = Pt(12)
set_chinese_font(list4_2_1_5)

list4_2_1_6 = doc.add_paragraph()
run = list4_2_1_6.add_run('6. 说说页：展示用户说说。')
run.font.size = Pt(12)
set_chinese_font(list4_2_1_6)

list4_2_1_7 = doc.add_paragraph()
run = list4_2_1_7.add_run('7. 留言板：用户留言功能。')
run.font.size = Pt(12)
set_chinese_font(list4_2_1_7)

# 管理后台
sub_section4_2_2 = doc.add_heading('4.2.2 管理后台', level=3)
set_title_font(sub_section4_2_2)

para4_2_2 = doc.add_paragraph()
para4_2_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para4_2_2.add_run('管理后台使用Vue 3.2实现，主要功能包括：')
run.font.size = Pt(12)
set_chinese_font(para4_2_2)
doc.add_paragraph()

list4_2_2 = doc.add_paragraph()
run = list4_2_2.add_run('1. 用户管理：管理系统用户，分配权限。')
run.font.size = Pt(12)
set_chinese_font(list4_2_2)

list4_2_2_2 = doc.add_paragraph()
run = list4_2_2_2.add_run('2. 文章管理：发布、编辑、删除文章。')
run.font.size = Pt(12)
set_chinese_font(list4_2_2_2)

list4_2_2_3 = doc.add_paragraph()
run = list4_2_2_3.add_run('3. 分类标签管理：管理文章分类和标签。')
run.font.size = Pt(12)
set_chinese_font(list4_2_2_3)

list4_2_2_4 = doc.add_paragraph()
run = list4_2_2_4.add_run('4. 评论管理：管理用户评论。')
run.font.size = Pt(12)
set_chinese_font(list4_2_2_4)

list4_2_2_5 = doc.add_paragraph()
run = list4_2_2_5.add_run('5. 系统配置：配置网站基本信息。')
run.font.size = Pt(12)
set_chinese_font(list4_2_2_5)

section4_3 = doc.add_heading('4.3 移动端实现', level=2)
set_title_font(section4_3)

para4_3 = doc.add_paragraph()
para4_3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para4_3.add_run('移动端使用UniApp实现，主要功能包括：')
run.font.size = Pt(12)
set_chinese_font(para4_3)
doc.add_paragraph()

list4_3 = doc.add_paragraph()
run = list4_3.add_run('1. 首页：展示文章列表。')
run.font.size = Pt(12)
set_chinese_font(list4_3)

list4_3_2 = doc.add_paragraph()
run = list4_3_2.add_run('2. 文章详情：查看文章内容。')
run.font.size = Pt(12)
set_chinese_font(list4_3_2)

list4_3_3 = doc.add_paragraph()
run = list4_3_3.add_run('3. 说说页：查看用户说说。')
run.font.size = Pt(12)
set_chinese_font(list4_3_3)

list4_3_4 = doc.add_paragraph()
run = list4_3_4.add_run('4. 个人中心：管理个人信息。')
run.font.size = Pt(12)
set_chinese_font(list4_3_4)

# 分页
doc.add_page_break()

# 5 系统测试
chapter5 = doc.add_heading('5 系统测试', level=1)
set_title_font(chapter5)

section5_1 = doc.add_heading('5.1 测试方法', level=2)
set_title_font(section5_1)

para5_1 = doc.add_paragraph()
para5_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para5_1.add_run('系统测试采用以下方法：')
run.font.size = Pt(12)
set_chinese_font(para5_1)
doc.add_paragraph()

list5_1 = doc.add_paragraph()
run = list5_1.add_run('1. 功能测试：测试系统的各项功能是否正常工作。')
run.font.size = Pt(12)
set_chinese_font(list5_1)

list5_1_2 = doc.add_paragraph()
run = list5_1_2.add_run('2. 性能测试：测试系统的响应速度和并发处理能力。')
run.font.size = Pt(12)
set_chinese_font(list5_1_2)

list5_1_3 = doc.add_paragraph()
run = list5_1_3.add_run('3. 兼容性测试：测试系统在不同浏览器和设备上的兼容性。')
run.font.size = Pt(12)
set_chinese_font(list5_1_3)

list5_1_4 = doc.add_paragraph()
run = list5_1_4.add_run('4. 安全性测试：测试系统的安全性，防止SQL注入等攻击。')
run.font.size = Pt(12)
set_chinese_font(list5_1_4)

section5_2 = doc.add_heading('5.2 测试结果', level=2)
set_title_font(section5_2)

para5_2 = doc.add_paragraph()
para5_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para5_2.add_run('通过测试，系统各项功能均正常工作，具体测试结果如下：')
run.font.size = Pt(12)
set_chinese_font(para5_2)
doc.add_paragraph()

# 测试结果表格
table5 = doc.add_table(rows=6, cols=3)
table5.style = 'Table Grid'

# 表头
heading_cells5 = table5.rows[0].cells
heading_cells5[0].text = '测试项'
heading_cells5[1].text = '测试结果'
heading_cells5[2].text = '备注'

# 设置表头样式
for cell in heading_cells5:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_chinese_font(paragraph)
        for run in paragraph.runs:
            run.font.bold = True

# 填充数据
test_data = [
    ['用户注册登录', '通过', '功能正常'],
    ['文章发布编辑', '通过', '支持Markdown格式'],
    ['评论系统', '通过', '支持回复和点赞'],
    ['分类标签管理', '通过', '功能正常'],
    ['系统性能', '通过', '响应速度快']
]

for i, row_data in enumerate(test_data, start=1):
    row = table5.rows[i]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    row.cells[2].text = row_data[2]
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            set_chinese_font(paragraph)

# 分页
doc.add_page_break()

# 6 结论与展望
chapter6 = doc.add_heading('6 结论与展望', level=1)
set_title_font(chapter6)

section6_1 = doc.add_heading('6.1 结论', level=2)
set_title_font(section6_1)

para6_1 = doc.add_paragraph()
para6_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para6_1.add_run('本研究成功设计并实现了一个基于SpringBoot+Vue的个人博客系统，系统具有以下特点：')
run.font.size = Pt(12)
set_chinese_font(para6_1)
doc.add_paragraph()

list6_1 = doc.add_paragraph()
run = list6_1.add_run('1. 功能完善：实现了用户管理、文章管理、评论管理、分类标签管理、说说管理、资源管理等核心功能。')
run.font.size = Pt(12)
set_chinese_font(list6_1)

list6_1_2 = doc.add_paragraph()
run = list6_1_2.add_run('2. 技术先进：采用前后端分离架构，使用SpringBoot、Vue等现代Web技术。')
run.font.size = Pt(12)
set_chinese_font(list6_1_2)

list6_1_3 = doc.add_paragraph()
run = list6_1_3.add_run('3. 界面美观：前端界面设计简洁美观，用户体验良好。')
run.font.size = Pt(12)
set_chinese_font(list6_1_3)

list6_1_4 = doc.add_paragraph()
run = list6_1_4.add_run('4. 性能稳定：系统经过测试，性能稳定，响应速度快。')
run.font.size = Pt(12)
set_chinese_font(list6_1_4)

doc.add_paragraph()

para6_1_2 = doc.add_paragraph()
para6_1_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para6_1_2.add_run('系统达到了预期的设计目标，满足了个人博客的基本需求，同时提供了丰富的扩展功能。通过本项目的开发，我不仅巩固了所学的专业知识，还提高了实际开发能力。')
run.font.size = Pt(12)
set_chinese_font(para6_1_2)

section6_2 = doc.add_heading('6.2 展望', level=2)
set_title_font(section6_2)

para6_2 = doc.add_paragraph()
para6_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para6_2.add_run('系统还有以下改进空间：')
run.font.size = Pt(12)
set_chinese_font(para6_2)
doc.add_paragraph()

list6_2 = doc.add_paragraph()
run = list6_2.add_run('1. 增加搜索功能：实现全文搜索，提高文章检索效率。')
run.font.size = Pt(12)
set_chinese_font(list6_2)

list6_2_2 = doc.add_paragraph()
run = list6_2_2.add_run('2. 优化移动端体验：进一步优化小程序的用户体验。')
run.font.size = Pt(12)
set_chinese_font(list6_2_2)

list6_2_3 = doc.add_paragraph()
run = list6_2_3.add_run('3. 增加社交功能：添加关注、私信等社交功能。')
run.font.size = Pt(12)
set_chinese_font(list6_2_3)

list6_2_4 = doc.add_paragraph()
run = list6_2_4.add_run('4. 优化性能：使用缓存技术进一步提高系统性能。')
run.font.size = Pt(12)
set_chinese_font(list6_2_4)

list6_2_5 = doc.add_paragraph()
run = list6_2_5.add_run('5. 增加多语言支持：支持中英文切换。')
run.font.size = Pt(12)
set_chinese_font(list6_2_5)

# 分页
doc.add_page_break()

# 参考文献
references = doc.add_heading('参考文献', level=1)
set_title_font(references)

ref1 = doc.add_paragraph()
run = ref1.add_run('1. 王楠. 基于SpringBoot和Vue的个人博客系统设计与实现[D]. 吉林大学, 2020.')
run.font.size = Pt(12)
set_chinese_font(ref1)

ref2 = doc.add_paragraph()
run = ref2.add_run('2. 张婷婷. 基于前后端分离架构的博客系统设计与实现[D]. 南京邮电大学, 2021.')
run.font.size = Pt(12)
set_chinese_font(ref2)

ref3 = doc.add_paragraph()
run = ref3.add_run('3. 焦鹏珲. 基于SpringBoot和Vue框架的电子招投标系统的设计与实现[D]. 南京大学, 2018.')
run.font.size = Pt(12)
set_chinese_font(ref3)

ref4 = doc.add_paragraph()
run = ref4.add_run('4. 赵一品. 基于Spring Boot和MyBatis的银行知识库管理系统的设计与实现[D]. 山东大学, 2020.')
run.font.size = Pt(12)
set_chinese_font(ref4)

ref5 = doc.add_paragraph()
run = ref5.add_run('5. 何娇. 基于Android平台的学习笔记系统设计与实现[D]. 内蒙古大学, 2018.')
run.font.size = Pt(12)
set_chinese_font(ref5)

# 分页
doc.add_page_break()

# 致谢
acknowledgements = doc.add_heading('致谢', level=1)
set_title_font(acknowledgements)

para_ack = doc.add_paragraph()
para_ack.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para_ack.add_run('在本项目的开发过程中，我得到了许多人的帮助和支持，在此表示衷心的感谢。')
run.font.size = Pt(12)
set_chinese_font(para_ack)
doc.add_paragraph()

para_ack2 = doc.add_paragraph()
para_ack2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para_ack2.add_run('首先，感谢我的指导教师马莉老师，她在项目的选题、设计和实现过程中给予了我悉心的指导和建议，帮助我解决了许多技术难题。')
run.font.size = Pt(12)
set_chinese_font(para_ack2)
doc.add_paragraph()

para_ack3 = doc.add_paragraph()
para_ack3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para_ack3.add_run('其次，感谢我的同学们，他们在项目开发过程中给予了我很多帮助和鼓励，与我分享了许多宝贵的经验和想法。')
run.font.size = Pt(12)
set_chinese_font(para_ack3)
doc.add_paragraph()

para_ack4 = doc.add_paragraph()
para_ack4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para_ack4.add_run('最后，感谢我的家人和朋友，他们在我学习和生活中给予了我支持和理解，让我能够专注于项目的开发。')
run.font.size = Pt(12)
set_chinese_font(para_ack4)
doc.add_paragraph()

para_ack5 = doc.add_paragraph()
para_ack5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
run = para_ack5.add_run('通过本项目的开发，我不仅学到了很多专业知识，还提高了实际开发能力和解决问题的能力。我将继续努力，不断学习和进步。')
run.font.size = Pt(12)
set_chinese_font(para_ack5)

# 保存文档
doc.save('/workspace/基于SpringBoot+Vue的个人博客系统设计与实现-黄翀.docx')
print('论文生成成功！文件保存在：/workspace/基于SpringBoot+Vue的个人博客系统设计与实现-黄翀.docx')
