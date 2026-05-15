from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


TEMPLATE_PATH = r"e:\桌面\shiyi-blog-master\2025-毕设设计（仅供参考）.docx"
MARKDOWN_PATH = r"e:\桌面\shiyi-blog-master\毕业论文-基于SpringBoot的个人博客系统-初稿.md"
OUTPUT_PATH = r"e:\桌面\shiyi-blog-master\基于SpringBoot的个人博客系统设计与实现-毕业论文.docx"

TITLE = "基于SpringBoot的个人博客系统设计与实现"
COLLEGE = "计算机科学与工程学院"
CLASS_NAME = "2022软件工程1班"
STUDENT_ID = "2022070030132"
STUDENT_NAME = "黄翀"
TEACHER = "马莉"
FINISH_DATE = "2026年6月"


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "右键更新域以生成目录"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)


def fill_cover_info(doc):
    for p in doc.paragraphs:
        t = p.text.strip()
        if p.style and p.style.name == "论文题目":
            p.text = TITLE
        elif t == "学    院：":
            p.text = f"学    院：{COLLEGE}"
        elif t == "专业班级：":
            p.text = f"专业班级：{CLASS_NAME}"
        elif t == "学    号：":
            p.text = f"学    号：{STUDENT_ID}"
        elif t == "学生姓名：":
            p.text = f"学生姓名：{STUDENT_NAME}"
        elif t == "指导教师：":
            p.text = f"指导教师：{TEACHER}"
        elif t.startswith("完成日期："):
            p.text = f"完成日期：{FINISH_DATE}"


def clear_from_abstract(doc):
    start_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "摘  要":
            start_idx = i
            break
    if start_idx is None:
        raise RuntimeError("模板中未找到“摘  要”段落。")

    for p in doc.paragraphs[start_idx:]:
        remove_paragraph(p)


def parse_markdown(md_text):
    lines = md_text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("---"):
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h0", line[2:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h1", line[3:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h2", line[4:].strip()))
            i += 1
            continue
        if line.startswith("#### "):
            blocks.append(("h3", line[5:].strip()))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(("table", table_lines))
            continue
        blocks.append(("p", line.strip()))
        i += 1
    return blocks


def add_markdown_table(doc, table_lines):
    rows = []
    for raw in table_lines:
        parts = [c.strip() for c in raw.strip().strip("|").split("|")]
        rows.append(parts)
    if len(rows) < 2:
        return
    if set(rows[1][0].replace("-", "").replace(":", "")) == {""}:
        data_rows = [rows[0]] + rows[2:]
    else:
        data_rows = rows
    cols = max(len(r) for r in data_rows)
    tbl = doc.add_table(rows=len(data_rows), cols=cols)
    tbl.style = "Table Grid"
    for r, row_data in enumerate(data_rows):
        for c in range(cols):
            txt = row_data[c] if c < len(row_data) else ""
            tbl.cell(r, c).text = txt


def write_content(doc, blocks):
    p = doc.add_paragraph("摘  要")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.style = "Normal"

    in_main = False
    started = False
    for kind, text in blocks:
        if kind == "h0":
            continue
        if kind != "table" and kind == "h1" and text == "摘要":
            started = True
            continue
        if not started:
            continue
        if kind != "table" and text == "摘要":
            continue
        if kind != "table" and text == "Abstract":
            p = doc.add_paragraph("ABSTRACT")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.style = "Normal"
            continue
        if kind != "table" and text.startswith("第1章 "):
            in_main = True
            doc.add_page_break()
        if kind != "table" and text == "第1章 绪论":
            doc.add_paragraph("目  录", style="toc 1")
            toc_p = doc.add_paragraph("", style="Normal")
            add_toc_field(toc_p)
            doc.add_page_break()

        if kind == "h1":
            if in_main:
                doc.add_paragraph(text, style="Heading 1")
            else:
                doc.add_paragraph(text, style="Normal")
        elif kind == "h2":
            doc.add_paragraph(text, style="Heading 2" if in_main else "Normal")
        elif kind == "h3":
            doc.add_paragraph(text, style="Heading 3" if in_main else "Normal")
        elif kind == "table":
            add_markdown_table(doc, text)
        else:
            para = doc.add_paragraph(text, style="Normal")
            para.paragraph_format.first_line_indent = doc.styles["Normal"].paragraph_format.first_line_indent


def main():
    doc = Document(TEMPLATE_PATH)
    fill_cover_info(doc)
    clear_from_abstract(doc)
    md_text = open(MARKDOWN_PATH, "r", encoding="utf-8").read()
    blocks = parse_markdown(md_text)
    write_content(doc, blocks)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
